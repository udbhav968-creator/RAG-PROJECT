import io
import logging
from fastapi import APIRouter, HTTPException, UploadFile, File, Form
from app.models import IngestRequest, IngestResponse, DocumentListResponse, DocumentItem
from app.core.rag_pipeline import rag_pipeline
from app.core.retrieval import get_all_documents, delete_document_chunks
from app.workers.tasks import ingest_document

logger = logging.getLogger(__name__)
router = APIRouter()

def _extract_file_text(file: UploadFile) -> str:
    filename = file.filename.lower()
    content_bytes = file.file.read()
    
    if filename.endswith(".pdf"):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            pages_text = [page.extract_text() for page in reader.pages if page.extract_text()]
            return "\n".join(pages_text)
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {e}")
            return content_bytes.decode('utf-8', errors='ignore')
            
    elif filename.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content_bytes))
            return "\n".join([p.text for p in doc.paragraphs if p.text])
        except Exception as e:
            logger.warning(f"docx extraction failed: {e}")
            return content_bytes.decode('utf-8', errors='ignore')
            
    else:
        return content_bytes.decode('utf-8', errors='ignore')

@router.post("/ingest", response_model=IngestResponse)
async def ingest(request: IngestRequest):
    try:
        # Try Celery task first
        try:
            task = ingest_document.delay(
                request.document_id,
                request.content,
                request.chunk_size or 500,
                request.chunk_overlap or 50
            )
            res = task.get(timeout=10)
            return IngestResponse(**res)
        except Exception as celery_err:
            logger.info(f"Celery task unavailable ({celery_err}). Ingesting inline.")
            res = rag_pipeline.ingest_document_text(
                document_id=request.document_id,
                content=request.content,
                chunk_size=request.chunk_size or 500,
                chunk_overlap=request.chunk_overlap or 50
            )
            return IngestResponse(**res)
    except Exception as e:
        logger.error(f"Ingestion failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))

from fastapi import APIRouter, HTTPException, Request

def _parse_bytes_text(filename: str, content_bytes: bytes) -> str:
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content_bytes))
            pages_text = [page.extract_text() for page in reader.pages if page.extract_text()]
            return "\n".join(pages_text)
        except Exception as e:
            logger.warning(f"pypdf extraction failed: {e}")
            return content_bytes.decode('utf-8', errors='ignore')
    elif filename_lower.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content_bytes))
            return "\n".join([p.text for p in doc.paragraphs if p.text])
        except Exception as e:
            logger.warning(f"docx extraction failed: {e}")
            return content_bytes.decode('utf-8', errors='ignore')
    else:
        return content_bytes.decode('utf-8', errors='ignore')

import re

def _parse_multipart_fallback(body_bytes: bytes):
    doc_id = None
    filename = "document.txt"
    file_bytes = body_bytes
    
    doc_match = re.search(rb'name="document_id"\r\n\r\n([^\r\n]+)', body_bytes)
    if doc_match:
        doc_id = doc_match.group(1).decode('utf-8', errors='ignore').strip()
        
    fn_match = re.search(rb'filename="([^"]+)"', body_bytes)
    if fn_match:
        filename = fn_match.group(1).decode('utf-8', errors='ignore').strip()
        if not doc_id:
            doc_id = filename.rsplit('.', 1)[0].upper()
            
    content_match = re.search(rb'filename="[^"]+"(?:\r\nContent-Type:[^\r\n]+)?\r\n\r\n(.*?)\r\n--', body_bytes, re.DOTALL)
    if content_match:
        file_bytes = content_match.group(1)
        
    return doc_id or "UPLOADED_DOC", filename, file_bytes

@router.post("/ingest/file", response_model=IngestResponse)
async def ingest_file(request: Request):
    try:
        content_type = request.headers.get("content-type", "")
        doc_id = request.headers.get("x-document-id")
        filename = request.headers.get("x-filename", "document.txt")
        content_bytes = b""

        if "multipart/form-data" in content_type:
            try:
                form = await request.form()
                file_obj = form.get("file")
                doc_id_form = form.get("document_id")
                if file_obj:
                    filename = getattr(file_obj, "filename", filename)
                    content_bytes = await file_obj.read()
                    if doc_id_form:
                        doc_id = str(doc_id_form)
                    else:
                        doc_id = filename.rsplit('.', 1)[0].upper()
                else:
                    body_bytes = await request.body()
                    fallback_id, filename, content_bytes = _parse_multipart_fallback(body_bytes)
                    doc_id = doc_id or fallback_id
            except Exception as form_err:
                logger.info(f"Form parsing fallback: {form_err}")
                body_bytes = await request.body()
                fallback_id, filename, content_bytes = _parse_multipart_fallback(body_bytes)
                doc_id = doc_id or fallback_id
        else:
            content_bytes = await request.body()
            doc_id = doc_id or "UPLOADED_DOC"

        extracted_text = _parse_bytes_text(filename, content_bytes)
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Unable to extract text content from file.")

        res = rag_pipeline.ingest_document_text(
            document_id=doc_id,
            content=extracted_text,
            chunk_size=500,
            chunk_overlap=50
        )
        return IngestResponse(**res)
    except Exception as e:
        logger.error(f"File ingestion failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))



@router.get("/documents", response_model=DocumentListResponse)
async def list_documents():
    docs_raw = get_all_documents()
    doc_items = [DocumentItem(**d) for d in docs_raw]
    return DocumentListResponse(
        total_documents=len(doc_items),
        documents=doc_items
    )

@router.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    removed = delete_document_chunks(doc_id)
    return {"status": "deleted", "document_id": doc_id, "chunks_removed": removed}